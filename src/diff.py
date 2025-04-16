import re
import json
import difflib
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH

KEYS = [
    "STUDY", "INDICATION", "COMPARISON", "ACCESSION NUMBER(S)", "ORDERING CLINICIAN",
    "TECHNIQUE", "FINDINGS", "IMPRESSION", "MACRO"
]

def load_json_data(file_path):
    """Load and parse the JSON data from a file."""
    with open(file_path, 'r', encoding='utf-8') as file:
        # Try to parse as a list
        content = file.read()
        try:
            # First try to parse as-is
            data = json.loads(content)
        except json.JSONDecodeError:
            # If it fails, try to add brackets to make it a proper list
            try:
                data = json.loads(f"[{content}]")
            except json.JSONDecodeError:
                # Last attempt: fix common JSON issues
                fixed_content = content.replace("'", '"')
                # Remove trailing commas
                fixed_content = re.sub(r',\s*}', '}', fixed_content)
                fixed_content = re.sub(r',\s*]', ']', fixed_content)
                try:
                    data = json.loads(f"[{fixed_content}]")
                except json.JSONDecodeError as e:
                    raise ValueError(f"Failed to parse JSON: {e}")
    return data

def parse_one_json_item(json_item):
    output = {}
    res_or_attending_key = list(json_item.keys())[0]

    text = json_item[res_or_attending_key]
    text = text.replace('\\r\\n', '\n').replace('\\n', '\n')
    text = text.replace('[', '').replace(']', '')
    # Remove excessive newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Clean up whitespace
    text = text.replace('.', '. ')
    text = re.sub(r'([a-zA-Z]):([a-zA-Z])', r'\1: \2', text)
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()

    relkeys = [key for key in KEYS if key in text]
    for idx, key in enumerate(relkeys):
        end_report_section_idx = text.index(relkeys[idx + 1]) if idx < len(relkeys)-1 else len(text)
        try:
            report_section = text[text.index(key) + len(key):end_report_section_idx].strip(": ").strip()
            output[key] = report_section
        except ValueError:
            output[key] = ""

    # Process the findings section into subsections based on the anatomy in particular
    findings_sections = re.split(r"([A-Z][A-Z\s]+:)", output["FINDINGS"])
    findings_dict = {}

    for subsection in findings_sections:  # skip first element which will be 'FINDINGS'
        current_subsection_header = None
        # Case where the element is a header
        if re.match(r"^[A-Z][A-Z\s]+:$", subsection):
            current_subsection_header = subsection.strip(": ").upper()
        # Case where the element is some content belonging to a header
        elif current_subsection_header is not None:
            findings_dict[current_subsection_header] = subsection.strip(": ").replace("\n", " ").strip()
        # Case where the text is the first text preceding any of the section headers
        else:
            findings_dict["GENERAL"] = subsection.strip(": ").replace("\n", " ").strip()

    output["FINDINGS"] = findings_dict
    return output

def preprocess_json(input_json):
    out = []
    for idx in range(0, len(input_json) - 2, 2):  # have to assume this goes by 2
        item1 = input_json[idx]
        item2 = input_json[idx+1]
        reader_keys = [list(item1.keys())[0], list(item2.keys())[0]]
        out.append([
            {"reader": reader_keys[0], **parse_one_json_item(item1)},
            {"reader": reader_keys[1], **parse_one_json_item(item2)}
        ])
    return out

def preprocess_text(text):
    """Perform text preprocessing on the input string."""
    # Replace escape sequences with their actual characters
    text = text.replace('\\r\\n', '\n').replace('\\n', '\n')
    # Remove excessive newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Clean up whitespace
    text = re.sub(r'\s+', ' ', text)
    # Restore paragraph breaks
    text = text.replace('. ', '.\n')
    return text

def get_intelligent_diff(text1, text2):
    """
    Generate an intelligent diff between two text strings.
    Returns a list of (text, format) tuples where format indicates
    whether text should be normal, strikethrough, or added.
    """
    # Preprocess both texts
    text1 = preprocess_text(text1)
    text2 = preprocess_text(text2)

    # If they're identical (after preprocessing), return early
    if text1 == text2:
        return [("(NO CORRECTIONS MADE)", "normal")]

    # Use SequenceMatcher for intelligent comparison
    matcher = difflib.SequenceMatcher(None, text1, text2)

    result = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            result.append((text1[i1:i2], "normal"))
        elif tag == 'delete':
            result.append((text1[i1:i2], "delete"))
        elif tag == 'insert':
            result.append((text2[j1:j2], "insert"))
        elif tag == 'replace':
            result.append((text1[i1:i2], "delete"))
            result.append((text2[j1:j2], "insert"))

    return result

def improve_diff_quality(text1, text2):
    """
    Improve diff quality by tokenizing and normalizing text.
    This is an alternative approach that can provide better results in some cases.
    """
    # Normalize whitespace and split into sentences
    def normalize_and_split(text):
        if not text:
            return []

        # Replace common abbreviations with placeholders to avoid period confusion
        abbr_map = {
            "Dr.": "DrPLACEHOLDER",
            "Mr.": "MrPLACEHOLDER",
            "Mrs.": "MrsPLACEHOLDER",
            "Ms.": "MsPLACEHOLDER",
            "etc.": "etcPLACEHOLDER",
            "i.e.": "iePLACEHOLDER",
            "e.g.": "egPLACEHOLDER",
            "vs.": "vsPLACEHOLDER",
            "Fig.": "FigPLACEHOLDER",
            "fig.": "figPLACEHOLDER",
            "Ref.": "RefPLACEHOLDER",
            "ref.": "refPLACEHOLDER",
            "cm.": "cmPLACEHOLDER",
            "mm.": "mmPLACEHOLDER",
            "a.m.": "amPLACEHOLDER",
            "p.m.": "pmPLACEHOLDER",
        }

        # Apply abbreviation replacements
        for abbr, placeholder in abbr_map.items():
            text = text.replace(abbr, placeholder)

        # Normalize spacing around punctuation
        text = re.sub(r'\s*([,.;:()])\s*', r'\1 ', text)

        # Split into sentences using regex
        sentences = re.split(r'(?<=[.!?])\s+', text)

        # Restore abbreviations
        for abbr, placeholder in abbr_map.items():
            sentences = [s.replace(placeholder, abbr) for s in sentences]

        return sentences

    # Split texts into sentences
    sentences1 = normalize_and_split(text1)
    sentences2 = normalize_and_split(text2)

    # Use SequenceMatcher for sentence-level diffing
    matcher = difflib.SequenceMatcher(None, sentences1, sentences2)

    result = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            result.append((" ".join(sentences1[i1:i2]), "normal"))
        elif tag == 'delete':
            result.append((" ".join(sentences1[i1:i2]), "delete"))
        elif tag == 'insert':
            result.append((" ".join(sentences2[j1:j2]), "insert"))
        elif tag == 'replace':
            # For replacements, try to do word-level diffing
            deleted_text = " ".join(sentences1[i1:i2])
            inserted_text = " ".join(sentences2[j1:j2])

            # If the texts are very different, just mark them as delete/insert
            if len(deleted_text) == 0 or len(inserted_text) == 0 or \
                abs(len(deleted_text) - len(inserted_text)) > 0.5 * max(len(deleted_text), len(inserted_text)):
                result.append((deleted_text, "delete"))
                result.append((inserted_text, "insert"))
            else:
                # Do word-level diffing
                words1 = re.findall(r'\S+|\s+', deleted_text)
                words2 = re.findall(r'\S+|\s+', inserted_text)

                word_matcher = difflib.SequenceMatcher(None, words1, words2)

                for w_tag, w_i1, w_i2, w_j1, w_j2 in word_matcher.get_opcodes():
                    if w_tag == 'equal':
                        result.append(("".join(words1[w_i1:w_i2]), "normal"))
                    elif w_tag == 'delete':
                        result.append(("".join(words1[w_i1:w_i2]), "delete"))
                    elif w_tag == 'insert':
                        result.append(("".join(words2[w_j1:w_j2]), "insert"))
                    elif w_tag == 'replace':
                        result.append(("".join(words1[w_i1:w_i2]), "delete"))
                        result.append(("".join(words2[w_j1:w_j2]), "insert"))

    # Check if result only contains whitespace changes
    def strip_whitespace(text):
        return re.sub(r'\s+', '', text)

    all_text = "".join([text for text, _ in result])
    if strip_whitespace(text1) == strip_whitespace(text2) and strip_whitespace(all_text) == strip_whitespace(text1):
        return [("(NO CORRECTIONS MADE)", "normal")]

    return result

def process_report_pair(resident, attending, doc):
    """Process a single resident-attending report pair and add to document."""
    # Add header
    header_text = f"STUDY: {resident.get('STUDY', '')}\n"
    header_text += f"INDICATION: {resident.get('INDICATION', '')}\n"
    header_text += f"ACCESSION NUMBER(S): {resident.get('ACCESSION NUMBER(S)', '')}"

    header = doc.add_heading(level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run(header_text)
    header_run.bold = True

    # Process FINDINGS section
    doc.add_heading("FINDINGS", level=2)

    # Extract and process FINDINGS content
    resident_findings = ""
    attending_findings = ""

    if 'FINDINGS' in resident:
        if isinstance(resident['FINDINGS'], dict):
            for section, text in resident['FINDINGS'].items():
                resident_findings += f"{text} "
        else:
            resident_findings = resident['FINDINGS']

    if 'FINDINGS' in attending:
        if isinstance(attending['FINDINGS'], dict):
            for section, text in attending['FINDINGS'].items():
                attending_findings += f"{text} "
        else:
            attending_findings = attending['FINDINGS']

    # Generate diff for FINDINGS
    findings_diff = improve_diff_quality(resident_findings, attending_findings)

    # Add FINDINGS diff to document
    if findings_diff == [("(NO CORRECTIONS MADE)", "normal")]:
        p = doc.add_paragraph()
        p.add_run("(NO CORRECTIONS MADE)")
    else:
        p = doc.add_paragraph()
        for text, format_type in findings_diff:
            run = p.add_run(text)
            if format_type == "delete":
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                run.font.strike = True
            elif format_type == "insert":
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

    # Process IMPRESSION section
    doc.add_heading("IMPRESSION", level=2)

    # Extract and process IMPRESSION content
    resident_impression = resident.get('IMPRESSION', '')
    attending_impression = attending.get('IMPRESSION', '')

    # Generate diff for IMPRESSION
    impression_diff = improve_diff_quality(resident_impression, attending_impression)

    # Add IMPRESSION diff to document
    if impression_diff == [("(NO CORRECTIONS MADE)", "normal")]:
        p = doc.add_paragraph()
        p.add_run("(NO CORRECTIONS MADE)")
    else:
        p = doc.add_paragraph()
        for text, format_type in impression_diff:
            run = p.add_run(text)
            if format_type == "delete":
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                run.font.strike = True
            elif format_type == "insert":
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN

def create_comparison_document_improved(data, output_file):
    """
    Create a Word document comparing resident and attending reads,
    using the improved diffing algorithm.
    """
    doc = Document()

    # Process each pair
    for i, pair in enumerate(data):
        # Make sure we have a resident and attending pair
        if len(pair) != 2:
            continue

        # Determine which is resident and which is attending
        resident_idx = 0 if pair[0].get('reader', '') == 'resident' else 1
        attending_idx = 1 - resident_idx

        if resident_idx >= len(pair) or attending_idx >= len(pair):
            continue

        resident = pair[resident_idx]
        attending = pair[attending_idx]

        # Process this report pair
        process_report_pair(resident, attending, doc)

        # Add page break after each pair except the last one
        if i < len(data) - 1:
            doc.add_page_break()

    # Save the document
    doc.save(output_file)
    print(f"Document saved to {output_file}")

def main():
    """Main function to run the program."""
    input_file = "output.json"
    output_file = "report_comparisons.docx"

    try:
        # Load data
        data = load_json_data(input_file)
        data = preprocess_json(data)

        # Create comparison document using improved algorithm
        create_comparison_document_improved(data, output_file)

        print(f"Successfully created comparison document: {output_file}")

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()
